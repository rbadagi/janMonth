#generates the resume
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# --- Header ---
name = doc.add_paragraph()
name.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = name.add_run("<Your Name>")
run.bold = True
run.font.size = Pt(22)

contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact.add_run("<Phone Number> | <Email Address> | <City, State> | LinkedIn: <LinkedIn URL>")

# --- Professional Summary ---
doc.add_heading("Professional Summary", level=1)
doc.add_paragraph(
    "Result-driven Automation Test Engineer with 8+ years of experience in manual and automation testing "
    "of web, API, and cloud-based applications. Proficient in designing robust test frameworks using Java and Python, "
    "with hands-on expertise in CI/CD pipelines, AWS services, and database validation. "
    "Strong analytical skills with a proven track record of improving software quality and reducing regression cycles."
)

# --- Technical Skills ---
doc.add_heading("Technical Skills", level=1)

skills = [
    ("Testing", "Manual Testing, Automation Testing, Functional Testing, Regression Testing, API Testing, Smoke & Sanity Testing"),
    ("Programming Languages", "Java, Python"),
    ("Automation Frameworks", "Selenium WebDriver, TestNG, Pytest, BDD (Cucumber), Page Object Model (POM)"),
    ("Cloud / AWS", "Step Functions, Glue Jobs, EC2, CloudWatch"),
    ("Database / Query", "SQL, Snowflake"),
    ("Version Control", "Git, GitHub, Bitbucket"),
    ("Web Technologies", "HTML, CSS (Basics)"),
    ("Tools", "Postman, ALM (HP Quality Center), Jira, Maven, Jenkins"),
    ("Methodologies", "Agile (Scrum), Waterfall, SDLC, STLC"),
]

table = doc.add_table(rows=len(skills), cols=2)
table.style = 'Light Grid Accent 1'
for i, (category, detail) in enumerate(skills):
    table.rows[i].cells[0].text = category
    table.rows[i].cells[1].text = detail

# --- Professional Experience ---
doc.add_heading("Professional Experience", level=1)

doc.add_heading("Project 1: <Project Name>", level=2)
p = doc.add_paragraph()
p.add_run("Client: ").bold = True
p.add_run("<Client Name>")
p = doc.add_paragraph()
p.add_run("Role: ").bold = True
p.add_run("Senior Automation Test Engineer")
p = doc.add_paragraph()
p.add_run("Duration: ").bold = True
p.add_run("<Start Date> - <End Date>")
p = doc.add_paragraph()
p.add_run("Environment: ").bold = True
p.add_run("Java, Selenium, TestNG, AWS (Step Functions, Glue, EC2, CloudWatch), Git, Jira, Snowflake")

doc.add_paragraph("Description:", style='List Bullet')
doc.add_paragraph("<Add project description here>", style='List Bullet 2')
doc.add_paragraph("Responsibilities:", style='List Bullet')
responsibilities_p1 = [
    "Designed and developed automation framework using Java, Selenium WebDriver, and TestNG.",
    "Validated data pipelines using AWS Glue jobs and monitored executions via CloudWatch.",
    "Performed API testing using Postman and automated API tests with Rest Assured.",
    "Wrote SQL queries for backend data validation in Snowflake.",
    "Monitored and triggered AWS Step Functions for orchestration testing.",
    "Managed test cases and defects in Jira and ALM.",
    "Participated in Agile ceremonies including sprint planning, daily standups, and retrospectives.",
]
for r in responsibilities_p1:
    doc.add_paragraph(r, style='List Bullet 2')

doc.add_paragraph()

doc.add_heading("Project 2: <Project Name>", level=2)
p = doc.add_paragraph()
p.add_run("Client: ").bold = True
p.add_run("<Client Name>")
p = doc.add_paragraph()
p.add_run("Role: ").bold = True
p.add_run("Automation Test Engineer")
p = doc.add_paragraph()
p.add_run("Duration: ").bold = True
p.add_run("<Start Date> - <End Date>")
p = doc.add_paragraph()
p.add_run("Environment: ").bold = True
p.add_run("Python, Pytest, Selenium, AWS EC2, SQL, Postman, Git, Jira")

doc.add_paragraph("Description:", style='List Bullet')
doc.add_paragraph("<Add project description here>", style='List Bullet 2')
doc.add_paragraph("Responsibilities:", style='List Bullet')
responsibilities_p2 = [
    "Developed automation scripts using Python and Pytest framework.",
    "Performed end-to-end testing of web applications deployed on AWS EC2 instances.",
    "Conducted manual and automated API testing using Postman.",
    "Executed SQL queries for data verification and validation.",
    "Collaborated with developers to identify root causes of defects and ensured timely resolution.",
    "Maintained test documentation and reported status in Jira.",
]
for r in responsibilities_p2:
    doc.add_paragraph(r, style='List Bullet 2')

doc.add_paragraph()

doc.add_heading("Project 3: <Project Name>", level=2)
p = doc.add_paragraph()
p.add_run("Client: ").bold = True
p.add_run("<Client Name>")
p = doc.add_paragraph()
p.add_run("Role: ").bold = True
p.add_run("QA Analyst / Test Engineer")
p = doc.add_paragraph()
p.add_run("Duration: ").bold = True
p.add_run("<Start Date> - <End Date>")
p = doc.add_paragraph()
p.add_run("Environment: ").bold = True
p.add_run("Java, Selenium, Maven, ALM, SQL, HTML/CSS, Git")

doc.add_paragraph("Description:", style='List Bullet')
doc.add_paragraph("<Add project description here>", style='List Bullet 2')
doc.add_paragraph("Responsibilities:", style='List Bullet')
responsibilities_p3 = [
    "Performed manual testing including functional, regression, and integration testing.",
    "Automated test cases using Java, Selenium WebDriver, and Maven.",
    "Managed test cases and execution cycles in HP ALM.",
    "Validated front-end UI elements using HTML/CSS knowledge.",
    "Performed database testing using SQL queries.",
    "Participated in requirement analysis and test planning activities.",
]
for r in responsibilities_p3:
    doc.add_paragraph(r, style='List Bullet 2')

# --- Education ---
doc.add_heading("Education", level=1)
doc.add_paragraph("<Degree> in <Major> - <University Name>, <Year of Graduation>")

# --- Certifications ---
doc.add_heading("Certifications", level=1)
doc.add_paragraph("<Add certifications here, e.g., ISTQB, AWS Cloud Practitioner, etc.>", style='List Bullet')

# Save
output_path = r"c:\Users\badagir\PycharmProjects\Automation\Automation_Test_Engineer_Resume.docx"
doc.save(output_path)
print(f"Resume saved at: {output_path}")
